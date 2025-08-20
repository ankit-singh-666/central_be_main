import os, re, json
from pathlib import Path
import pandas as pd

from llama_index.core import Document, VectorStoreIndex, Settings, StorageContext, SimpleDirectoryReader
from llama_index.embeddings.openai import OpenAIEmbedding
from llama_index.vector_stores.chroma import ChromaVectorStore
from llama_index.readers.google import GoogleDriveReader
from llama_index.core.vector_stores import MetadataFilter, MetadataFilters, FilterOperator, FilterCondition
from docx import Document as DocxDocument
from googleapiclient.discovery import build
from google.oauth2 import service_account
from googleapiclient.http import MediaIoBaseDownload
import io
import re
from openai import OpenAI
import chromadb
import tempfile
import zipfile
from parser import process_document_enhanced
import config


class MultiSourceIndexer:
    def __init__(
        self,
        chroma_path: str = "./embeddings_storage/chroma_db",
        embedding_model: str = config.OPENAI_MODELS["embedding"],
        openai_api_key: str = config.OPENAI_API_KEY
    ):
        self.embedding_model = embedding_model
        self.chroma_path = chroma_path
        self.index = None
        self.openai_api_key = openai_api_key or os.getenv("OPENAI_API_KEY")

        if not self.openai_api_key:
            raise ValueError("OpenAI API key is required. Set it via openai_api_key or environment variable.")

        self._setup_embedding()

    def _setup_embedding(self):
        Settings.embed_model = OpenAIEmbedding(
            model=self.embedding_model,
            api_key=self.openai_api_key
        )
        Settings.chunk_size = 512
        Settings.chunk_overlap = 50

    def _get_processors(self):
        def process_text(path: Path):
            text = path.read_text(encoding="utf-8", errors="ignore")
            return [Document(text=text, metadata={"file_name": path.name})] if text.strip() else []

        def process_csv(path: Path):
            df = pd.read_csv(path)
            txt = f"CSV: {path.name}\nColumns: {', '.join(df.columns)}\n{df.to_string(index=False)}"
            return [Document(text=txt, metadata={"file_name": path.name})]

        def process_json(path: Path):
            data = json.loads(path.read_text(encoding="utf-8"))
            return [Document(text=json.dumps(data, indent=2), metadata={"file_name": path.name})]

        def process_pdf(path: Path):
            docs = SimpleDirectoryReader(input_files=[str(path)]).load_data()
            for d in docs:
                d.metadata["file_name"] = path.name
            return docs

        def process_docx(path: Path):
            doc = DocxDocument(path)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            return [Document(text=text, metadata={"file_name": path.name})] if text.strip() else []

        return {
            ".txt": process_text,
            ".md": process_text,
            ".csv": process_csv,
            ".json": process_json,
            ".pdf": process_pdf,
            ".docx": process_docx
        }

    def build_index_from_file(self, file_path: str, collection_name: str = "all_files"):
        path = Path(file_path)
        if not path.exists() or not path.is_file():
            raise ValueError(f"File not found: {file_path}")

        processors = self._get_processors()
        suffix = path.suffix.lower()
        if suffix not in processors:
            raise ValueError(f"Unsupported file format: {suffix}")

        docs = processors[suffix](path)
        if not docs:
            raise ValueError("No documents extracted from file.")
        
        for doc in docs:
            doc.metadata["file_type"] = suffix
            doc.metadata["file_name"] = path.name
            doc.metadata["source"] = "local_file"

        self.index = self._create_index(docs, collection_name)
        print(f"✅ Single file index built from {path.name} and ready.")

    
   
   
   
   
   
   
   
   
    
    def build_index_from_drive(self, drive_link: str, collection_name: str = "all_files"):
        # Extract folder ID from the link
        match = re.search(r'/folders/([a-zA-Z0-9_-]+)', drive_link)
        if not match:
            raise ValueError("Invalid Google Drive folder link.")
        folder_id = match.group(1)
        print(f"📥 Loading files from Drive folder {folder_id}...")

        # Load service account credentials
        creds = service_account.Credentials.from_service_account_file("service_key.json")
        drive_service = build("drive", "v3", credentials=creds)

        # List all files in the folder
        results = drive_service.files().list(
            q=f"'{folder_id}' in parents and trashed=false",
            fields="files(id, name, mimeType)"
        ).execute()
        files = results.get("files", [])
        print(f"✅ Found {len(files)} files in Drive folder.")

        if not files:
            raise ValueError("No files found in Drive folder.")

        enriched_docs = []

        for f in files:
            file_id = f["id"]
            file_name = f["name"]
            file_ext = Path(file_name).suffix.lower()

            # Only enrich PDFs and DOCX files
            if file_ext not in [".pdf", ".docx"]:
                continue

            try:
                print(f"✨ Downloading {file_name} from Drive...")
                request = drive_service.files().get_media(fileId=file_id)
                fh = io.BytesIO()
                downloader = MediaIoBaseDownload(fh, request)
                done = False
                while not done:
                    status, done = downloader.next_chunk()

                # Save to temporary file
                temp_input = Path(tempfile.gettempdir()) / file_name
                with open(temp_input, "wb") as tmpf:
                    tmpf.write(fh.getvalue())

                print(f"✨ Enriching {file_name} with process_document_enhanced...")
                enriched_md_path = process_document_enhanced(str(temp_input))

                # Load enriched markdown
                md_docs = SimpleDirectoryReader(input_files=[enriched_md_path]).load_data()
                for d in md_docs:
                    d.metadata["file_name"] = file_name
                    d.metadata["file_type"] = file_ext
                    d.metadata["source"] = "drive_files"

                enriched_docs.extend(md_docs)

            except Exception as e:
                print(f"⚠️ Error processing {file_name}: {e}, skipping...")

        if not enriched_docs:
            raise ValueError("No documents ready after enrichment from Drive folder.")

        # Create the index
        self.index = self._create_index(enriched_docs, collection_name)
        print("✅ Drive folder index built (with enrichment) and ready.")
        
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    def build_index_from_zip(self, zip_path: str, collection_name: str = "zip_files"):
        path = Path(zip_path)
        if not path.exists() or not path.is_file():
            raise ValueError(f"ZIP file not found: {zip_path}")
        if path.suffix.lower() != ".zip":
            raise ValueError(f"Unsupported file format for ZIP indexing: {path.suffix}")

        processors = self._get_processors()
        docs = []

        with tempfile.TemporaryDirectory() as tmpdir:
            with zipfile.ZipFile(path, "r") as zip_ref:
                zip_ref.extractall(tmpdir)

            tmpdir_path = Path(tmpdir)
            for f in tmpdir_path.rglob("*"):
                if f.is_file() and f.suffix.lower() in processors:
                    try:
                        docs.extend(processors[f.suffix.lower()](f))
                        print(f"Processed from ZIP: {f.name}")
                    except Exception as e:
                        print(f"⚠️ Error processing {f.name}: {e}")

        if not docs:
            raise ValueError("No supported files found inside the ZIP archive.")

        self.index = self._create_index(docs, collection_name)
        print(f"✅ ZIP archive index built from {path.name} and ready.")

    def _create_index(self, docs, collection_name):
        chroma_client = chromadb.PersistentClient(path=self.chroma_path)
        vector_store = ChromaVectorStore(
            chroma_collection=chroma_client.get_or_create_collection(collection_name)
        )
        storage_context = StorageContext.from_defaults(vector_store=vector_store)
        return VectorStoreIndex.from_documents(docs, storage_context=storage_context, show_progress=True)

    def query(self, query_text: str, k: int = 3, source: str = None, file_type: str = None):
        if not self.index:
            raise RuntimeError("No index built yet. Use one of the build_index_* methods first.")
        
        filters_list = []
        if source:
            filters_list.append(MetadataFilter(key="source", value=source, operator=FilterOperator.EQ))
        if file_type:
            filters_list.append(MetadataFilter(key="file_type", value=file_type, operator=FilterOperator.EQ))

        metadata_filters = MetadataFilters(
            filters=filters_list,
            condition=FilterCondition.AND
        ) if filters_list else None

        retriever = self.index.as_retriever(similarity_top_k=k, filters=metadata_filters)
        results = retriever.retrieve(query_text)
            
        context_texts = []
        for r in results:
            fname = r.metadata.get("file_name", "Unknown")
            chunk = r.node.get_content()
            context_texts.append(f"[File: {fname}]\n{chunk}")
        context = "\n\n".join(context_texts)
        return context


if __name__ == "__main__":
    indexer = MultiSourceIndexer()